"""Render _case_summary_<chart>.md to a clean Word .docx for review.

Handles:
- # / ## / ### headings
- markdown pipe tables (| col | col |)
- bullets (- item)
- bold (**text**)
- blockquotes (> ...)
- inline code `x`

Output: _filled_docx/<chart>_<姓O名>_review.docx
"""
import io
import re
import sys
from pathlib import Path

# Force UTF-8 stdout so Chinese paths and unicode marks don't choke on cp950
if sys.stdout.encoding != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
SUMMARY_GLOB = "_case_summary_*.md"
OUT_DIR = ROOT / "_filled_docx"
OUT_DIR.mkdir(exist_ok=True)


# ---------------- Style helpers ----------------

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_runs_with_bold(paragraph, text):
    """Split text on **bold** markers and add runs accordingly."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # also handle `code` inline
            sub_parts = re.split(r"(`[^`]+`)", part)
            for sp in sub_parts:
                if sp.startswith("`") and sp.endswith("`"):
                    run = paragraph.add_run(sp[1:-1])
                    run.font.name = "Consolas"
                else:
                    paragraph.add_run(sp)


def setup_styles(doc):
    """Set default font + heading sizes for Chinese-friendly output."""
    style = doc.styles["Normal"]
    style.font.name = "Microsoft JhengHei"
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    for level, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        s = doc.styles[level]
        s.font.name = "Microsoft JhengHei"
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        s.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


# ---------------- Markdown parser ----------------

def parse_table_block(lines, start):
    """Parse a markdown pipe table starting at lines[start]; return (rows, end_idx).
    Skip the |---|---| separator row.
    """
    rows = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        line = lines[i].strip()
        # strip leading/trailing |
        cells = [c.strip() for c in line.strip("|").split("|")]
        # skip the alignment separator (|---|:---:|---:|)
        if all(re.match(r"^:?-+:?$", c) for c in cells):
            i += 1
            continue
        rows.append(cells)
        i += 1
    return rows, i


def render_table(doc, rows):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    # autofit
    table.autofit = True

    for r_idx, row_cells in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx in range(n_cols):
            cell = row.cells[c_idx]
            cell.text = ""  # clear
            text = row_cells[c_idx] if c_idx < len(row_cells) else ""
            p = cell.paragraphs[0]
            add_runs_with_bold(p, text)
            # Header row shading
            if r_idx == 0:
                set_cell_shading(cell, "D9E2F3")
                for run in p.runs:
                    run.bold = True


def render_md(doc, md_text):
    """Stream markdown line-by-line into docx."""
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # blank line
        if not stripped:
            i += 1
            continue

        # ATX headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        # horizontal rule
        if re.match(r"^-{3,}$", stripped):
            i += 1
            continue

        # blockquote
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run("⚠ ")
            run.bold = True
            add_runs_with_bold(p, stripped[2:])
            for run in p.runs:
                if run.text.startswith("⚠"):
                    continue
                run.italic = True
            i += 1
            continue

        # bulleted list
        if re.match(r"^\s*[-*]\s", line):
            p = doc.add_paragraph(style="List Bullet")
            text = re.sub(r"^\s*[-*]\s+", "", line)
            # Handle [x] / [ ] checkboxes
            cb_m = re.match(r"^\[([x ])\]\s+(.*)", text)
            if cb_m:
                box = "✓ " if cb_m.group(1) == "x" else "☐ "
                run = p.add_run(box)
                run.bold = (cb_m.group(1) == "x")
                add_runs_with_bold(p, cb_m.group(2))
            else:
                add_runs_with_bold(p, text)
            i += 1
            continue

        # Pipe table
        if stripped.startswith("|"):
            rows, end_i = parse_table_block(lines, i)
            render_table(doc, rows)
            i = end_i
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_runs_with_bold(p, stripped)
        i += 1


# ---------------- Main ----------------

def derive_output_name(chart, md_text):
    """Try to extract patient name from md; return safe filename."""
    name_m = re.search(r"\|\s*姓名\s*\|\s*([^（|（]+)", md_text)
    if name_m:
        name = name_m.group(1).strip()
        # mask middle character: 王小明 -> 王O明
        if len(name) >= 3:
            masked = name[0] + "O" + name[-1]
        elif len(name) == 2:
            masked = name[0] + "O"
        else:
            masked = name
        return f"{chart}_{masked}_review.docx"
    return f"{chart}_review.docx"


def main():
    md_files = sorted(ROOT.glob(SUMMARY_GLOB))
    if not md_files:
        raise SystemExit(f"找不到 {SUMMARY_GLOB} — 先跑完 EMR 抓取流程")

    for md_path in md_files:
        chart_m = re.match(r"_case_summary_(\d+)", md_path.stem)
        chart = chart_m.group(1) if chart_m else md_path.stem

        md_text = md_path.read_text(encoding="utf-8")
        out_name = derive_output_name(chart, md_text)
        out_path = OUT_DIR / out_name

        print(f"[render] {md_path.name} -> {out_path}")

        doc = Document()
        setup_styles(doc)

        # Cover banner (small)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("評鑑案例整理（病歷號 " + chart + "）")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        render_md(doc, md_text)

        doc.save(str(out_path))
        print(f"   ✓ {out_path.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
