"""Render a filled PCI cert cheatsheet docx from the template + a filled YAML.

Approach:
- Open the template docx with python-docx
- Walk the tables; each is a 2-column row: [label | value/example]
- For each label we recognize, write the value into the right cell, replacing any
  example placeholder text (anything matching `(例:...)` or empty).
- For the Evidence Checklist table, replace ☐ with ☑ where `found: true`, keep ☐
  for false/null, and append the location/note into the rightmost cell.
- Save as `out_docx/<chart>_<name_code>_filled.docx`.
"""
import sys
import io
import re
import yaml
from pathlib import Path
from copy import deepcopy

if sys.stdout.encoding != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document  # type: ignore
from docx.oxml.ns import qn

from config import (
    TEMPLATE_DOCX as TEMPLATE, FILLED_DOCX_DIR as OUT_DIR,
    GEMINI_WORKSPACE, load_cases,
)

OUT_DIR.mkdir(exist_ok=True)


def cell_text(cell):
    return "".join(p.text for p in cell.paragraphs)


def set_cell_text(cell, text, source=None):
    """Replace cell's text with `text`. If `source` given, append a
    parenthesized "[來源: ...]" line in italics underneath."""
    p0 = cell.paragraphs[0]
    for p in list(cell.paragraphs[1:]):
        p._element.getparent().remove(p._element)
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)
    lines = (text or "").split("\n")
    p0.add_run(lines[0])
    for line in lines[1:]:
        new_p = cell.add_paragraph()
        new_p.add_run(line)
    if source:
        src_p = cell.add_paragraph()
        run = src_p.add_run(f"[來源: {source}]")
        run.italic = True
        run.font.size = None  # let inheritance reduce; no hard size to avoid styling fights


def label_match(label, target):
    """Match if `target` keyword appears in cleaned label text."""
    return target in label


# ---------- Section ① — ⑥ scalar fillers ----------

# Map (label keyword) -> function(yaml) -> str
def scalar_fillers(y):
    """Return list of (label_keyword, value) tuples for the 2-col scalar tables."""
    stents = y.get("stents") or []
    stent_lines = []
    for s in stents:
        site = s.get("site", "?")
        typ = s.get("type", "?")
        bs = s.get("brand_size") or ""
        cnt = s.get("count", "?")
        stent_lines.append(f"{site}: {typ} × {cnt}" + (f"  [{bs}]" if bs else ""))
    stent_block = "\n".join(stent_lines) if stent_lines else "—"

    total_count = y.get("total_stent_count_this_session")
    prior_count = y.get("prior_stent_count")
    total_str = f"本次 {total_count} 支"
    if prior_count:
        total_str += f"  (含 prior {prior_count} 支 → 累積 {(total_count or 0)+prior_count})"

    timi = y.get("final_timi_flow") or "—"
    contrast = y.get("contrast_volume_ml")
    contrast_str = f"{contrast} mL" if contrast is not None else "—"
    fluoro = y.get("fluoroscopy_time_min")
    fluoro_str = f"{fluoro} min" if fluoro is not None else "—"
    door_to_dev = y.get("door_to_device_min")
    door_str = f"{door_to_dev} min" if door_to_dev is not None else "—"

    syntax = y.get("syntax_score")
    syntax_str = str(syntax) if syntax not in (None, "null") else "—"

    # (label_keyword, value, source)
    fillers = [
        ("病歷號 Chart No.", y.get("chart_no", ""), None),
        ("姓名", y.get("name_code", ""), None),
        ("PCI 日期", y.get("pci_date", ""), None),
        ("Operator", y.get("operator", ""), y.get("operator_src")),
        ("Demographics", y.get("demographics_comorbidities", ""),
            y.get("demographics_comorbidities_src")),
        ("Presentation", y.get("presentation", ""), y.get("presentation_src")),
        ("病灶位置", y.get("lesion_summary", ""), y.get("lesion_summary_src")),
        ("SYNTAX score", syntax_str, y.get("syntax_score_src")),
        ("為何選 PCI", y.get("why_pci_not_cabg", ""), y.get("why_pci_not_cabg_src")),
        ("Stent 數量", stent_block + "\n\n" + total_str, y.get("stents_src")),
        ("Total stent length",
            f"{y.get('total_stent_length_mm')} mm" if y.get("total_stent_length_mm") is not None else "—",
            y.get("stents_src")),
        ("使用的 imaging", y.get("imaging_used", ""), y.get("imaging_used_src")),
        ("Adjunctive devices", y.get("adjunctive_devices", ""), y.get("adjunctive_devices_src")),
        ("為何需要 >5 支", y.get("why_more_than_5_stents", ""), None),
        ("Final TIMI flow", timi, y.get("final_timi_flow_src")),
        ("Procedural complication", y.get("procedural_complication", ""),
            y.get("procedural_complication_src")),
        ("Contrast volume", contrast_str, y.get("contrast_volume_src")),
        ("Fluoroscopy time", fluoro_str, y.get("fluoroscopy_time_src")),
        ("Door-to-device", door_str, y.get("door_to_device_src")),
        ("被點到時這樣開場", y.get("opening_one_liner", ""), None),
    ]
    return fillers


# ---------- Evidence Checklist (section ⑤) ----------

CHECKLIST_LABELS = [
    ("PCI 知情同意書", "pci_consent"),
    ("Heart Team", "heart_team_or_cath_conference"),
    ("CV Surgery 照會", "cv_surgery_consult"),
    ("其他相關照會", "other_consults"),
    ("家庭會議", "family_meeting"),
    ("CDS 警示", "cds_alerts"),
    ("術前 echo", "preop_imaging_or_function"),
    ("術後追蹤", "postop_followup"),
]


def fill_checklist(table, checklist):
    """The checklist table has rows: [☐, 項目, 病歷位置/備註]. Fill rows by matching label."""
    if not checklist:
        return
    for row in table.rows:
        if len(row.cells) < 3:
            continue
        item_text = cell_text(row.cells[1])
        for label_key, yaml_key in CHECKLIST_LABELS:
            if label_key in item_text:
                entry = checklist.get(yaml_key) or {}
                found = entry.get("found")
                checkbox = "☑" if found is True else ("☒" if found is False else "☐")
                set_cell_text(row.cells[0], checkbox)
                loc = entry.get("location") or ""
                note = entry.get("note") or ""
                merged = []
                if loc:
                    merged.append(f"位置/日期: {loc}")
                if note:
                    merged.append(f"備註: {note}")
                if merged:
                    set_cell_text(row.cells[2], "\n".join(merged))
                break


# ---------- Q&A (section ⑦) ----------

def fill_qa(doc, questions):
    """Find the Q1/Q2/Q3 rows and fill each. Q labels in template are 'Q1:' etc."""
    q_idx = 0
    for tbl in doc.tables:
        for row in tbl.rows:
            if len(row.cells) < 2:
                continue
            label = cell_text(row.cells[0]).strip()
            m = re.match(r"^Q(\d+):", label)
            if m:
                if q_idx < len(questions):
                    q = questions[q_idx]
                    set_cell_text(row.cells[0], f"Q{m.group(1)}: {q.get('q', '')}")
                    set_cell_text(row.cells[1], q.get("a", ""))
                q_idx += 1


# ---------- Main ----------

GROUP2_TITLE_RE = re.compile(r"PCI\s*認證病歷小抄\s*—\s*第二組")


def drop_group_two(doc):
    """The template has both 第一組 (Stent>5) and 第二組 (Cover stent) — delete
    everything from the Group-2 title table onward."""
    body = doc.element.body
    boundary = None
    for child in list(body):
        text = "".join(t.text or "" for t in child.iter(qn("w:t")))
        if GROUP2_TITLE_RE.search(text):
            boundary = child
            break
    if boundary is None:
        return
    seen = False
    for child in list(body):
        if child is boundary:
            seen = True
        if seen:
            if child.tag == qn("w:sectPr"):
                continue
            body.remove(child)


def drop_group_one(doc):
    """For group=2 cases: delete everything BEFORE the 第二組 title table so the
    output starts with the Group-2 cheatsheet."""
    body = doc.element.body
    boundary = None
    for child in list(body):
        text = "".join(t.text or "" for t in child.iter(qn("w:t")))
        if GROUP2_TITLE_RE.search(text):
            boundary = child
            break
    if boundary is None:
        return
    for child in list(body):
        if child is boundary:
            return
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


# ---------- Group 2 fillers ----------

def _fmt_cover_stent(cs):
    if not cs:
        return "—"
    if cs.get("used") is False:
        return "未使用"
    if cs.get("used") is None:
        return "—"
    bs = cs.get("brand_size") or "?"
    pos = cs.get("position") or "?"
    return f"使用:{bs};部署位置:{pos}"


def _fmt_pericardiocentesis(p):
    if not p:
        return "—"
    if p.get("performed") is False:
        return "未執行"
    if p.get("performed") is None:
        return "—"
    parts = ["執行"]
    if p.get("drained_ml") is not None:
        parts.append(f"引流量 {p['drained_ml']} mL")
    if p.get("time"):
        parts.append(f"時間 {p['time']}")
    return ";".join(parts)


def _fmt_cv_surgery_standby(s):
    if not s:
        return "—"
    if s.get("activated") is False:
        return "未啟動"
    if s.get("activated") is None:
        return "—"
    parts = ["啟動"]
    if s.get("notify_time"):
        parts.append(f"通知 {s['notify_time']}")
    if s.get("arrival_time"):
        parts.append(f"到場 {s['arrival_time']}")
    return ";".join(parts)


def _fmt_mcs(m):
    if not m:
        return "—"
    devs = []
    if m.get("ecmo"):
        devs.append("ECMO")
    if m.get("iabp"):
        devs.append("IABP")
    if not devs:
        return "未使用"
    txt = "+".join(devs)
    if m.get("device_start_time"):
        txt += f" (啟動時間 {m['device_start_time']})"
    return txt


def _fmt_subtype(s):
    cover = "☑" if s == "cover_stent" else "☐"
    elective = "☑" if s == "elective_with_complication" else "☐"
    return f"{cover} Cover stent  {elective} Elective PCI with complication"


def _fmt_dapt(d):
    if not d:
        return "—"
    drugs = d.get("drugs") or "—"
    dur = d.get("duration") or "—"
    return f"{drugs}\n期程:{dur}"


def _fmt_status(s):
    return {
        "alive_stable": "Alive & stable",
        "needs_rehab": "需 rehab",
        "transfer_out": "轉院",
        "expired": "Expired",
    }.get(s, s or "—")


def scalar_fillers_g2(y):
    """Return [(label_keyword, value, src)] for Group-2 scalar tables."""
    cs = y.get("cover_stent_used") or {}
    pc = y.get("pericardiocentesis") or {}
    cv = y.get("cv_surgery_standby") or {}
    mcs = y.get("mcs_devices") or {}

    los = y.get("length_of_stay_days")
    los_str = f"{los} 天" if los is not None else "—"

    lvef = y.get("final_lvef")
    lvef_str = f"{lvef} %" if lvef is not None else "—"

    return [
        ("病歷號 Chart No.", y.get("chart_no", ""), None),
        ("姓名", y.get("name_code", ""), None),
        ("PCI 日期", y.get("pci_date", ""), None),
        ("Operator", y.get("operator", ""), y.get("operator_src")),
        ("本組分類", _fmt_subtype(y.get("group2_subtype")), y.get("group2_subtype_src")),
        ("Demographics", y.get("demographics_comorbidities", ""),
            y.get("demographics_comorbidities_src")),
        ("Presentation", y.get("presentation", ""), y.get("presentation_src")),
        ("原計畫 PCI 策略", y.get("original_pci_plan", ""), y.get("original_pci_plan_src")),
        ("Complication 類別", y.get("complication_category", ""),
            y.get("complication_category_src")),
        ("發生機轉", y.get("complication_mechanism", ""),
            y.get("complication_mechanism_src")),
        ("是否使用 cover stent", _fmt_cover_stent(cs), y.get("cover_stent_src")),
        ("是否需 pericardiocentesis", _fmt_pericardiocentesis(pc),
            y.get("pericardiocentesis_src")),
        ("是否啟動 CV surgery stand-by", _fmt_cv_surgery_standby(cv),
            y.get("cv_surgery_standby_src")),
        ("是否需 ECMO / IABP", _fmt_mcs(mcs), y.get("mcs_devices_src")),
        ("住院天數", los_str, y.get("length_of_stay_src")),
        ("出院時 status", _fmt_status(y.get("discharge_status")),
            y.get("discharge_status_src")),
        ("Final LVEF", lvef_str, y.get("final_lvef_src")),
        ("DAPT 計畫", _fmt_dapt(y.get("dapt_plan")), y.get("dapt_plan_src")),
        ("M&M 結論", y.get("mm_takeaway", ""), y.get("mm_takeaway_src")),
        ("被點到時這樣開場", y.get("opening_one_liner", ""), None),
    ]


# Group-2 timeline (table with 8 rows: header + 7 events)
TIMELINE_KEYS = [
    ("併發症發生", "complication_onset"),
    ("Bailout 決策", "bailout_decision"),
    ("Cover stent 部署", "cover_stent_deployment"),
    ("CV surgery 通知", "cv_surgery_notify"),
    ("家屬告知", "family_notification"),
    ("送 ICU / CCU", "icu_ccu_transfer"),
    ("後續處置", "followup_management"),
]


def fill_timeline_g2(table, timeline):
    if not timeline:
        return
    for row in table.rows:
        if len(row.cells) < 3:
            continue
        event_text = cell_text(row.cells[1])
        for label, key in TIMELINE_KEYS:
            if label in event_text:
                entry = timeline.get(key) or {}
                t = entry.get("time") or "—"
                set_cell_text(row.cells[0], t)
                note = entry.get("note") or ""
                if note:
                    set_cell_text(row.cells[1], f"{label}\n{note}")
                src = entry.get("src") or ""
                if src:
                    set_cell_text(row.cells[2], src)
                break


# Group-2 evidence checklist (11 rows)
CHECKLIST_LABELS_G2 = [
    ("原 PCI 知情同意書", "original_pci_consent"),
    ("Cover stent / Bailout 同意書", "cover_stent_or_bailout_consent"),
    ("併發症當下家屬告知", "family_notification_record"),
    ("CV Surgery 即時照會", "cv_surgery_realtime_consult"),
    ("Anesthesia / ICU / CCU", "anesthesia_icu_ccu_consult"),
    ("家庭會議", "family_meeting_followup"),
    ("Cath conference", "cath_conference_or_mm"),
    ("院內不良事件通報", "hospital_incident_report"),
    ("CDS", "cds_alerts"),
    ("後續追蹤", "postop_followup"),
]


def fill_checklist_g2(table, checklist):
    if not checklist:
        return
    for row in table.rows:
        if len(row.cells) < 3:
            continue
        item_text = cell_text(row.cells[1])
        for label_key, yaml_key in CHECKLIST_LABELS_G2:
            if label_key in item_text:
                entry = checklist.get(yaml_key) or {}
                found = entry.get("found")
                checkbox = "☑" if found is True else ("☒" if found is False else "☐")
                set_cell_text(row.cells[0], checkbox)
                loc = entry.get("location") or ""
                note = entry.get("note") or ""
                merged = []
                if loc:
                    merged.append(f"位置/日期: {loc}")
                if note:
                    merged.append(f"備註: {note}")
                if merged:
                    set_cell_text(row.cells[2], "\n".join(merged))
                break


def fill_qa_g2(doc, questions):
    """Group 2 has 4 pre-filled questions; we only fill the answer cell.
    Match by substring of the existing question text."""
    if not questions:
        return
    for tbl in doc.tables:
        for row in tbl.rows:
            if len(row.cells) < 2:
                continue
            label = cell_text(row.cells[0]).strip()
            if not label.startswith("Q:"):
                continue
            for q in questions:
                qtext = (q.get("q") or "").strip()
                if not qtext:
                    continue
                # match by first 6 meaningful chars
                core = re.sub(r"[Q:?? ]", "", qtext)[:6]
                if core and core in label.replace("?", "").replace("?", ""):
                    set_cell_text(row.cells[1], q.get("a", ""))
                    break


def fill_doc_g2(template_path, yaml_data, output_path):
    """Group 2 render path: drop group 1, fill group 2."""
    doc = Document(str(template_path))
    drop_group_one(doc)

    fillers = scalar_fillers_g2(yaml_data)
    used = set()

    for tbl in doc.tables:
        for row in tbl.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            label = cell_text(cells[0]).strip()
            if not label:
                continue
            for i, (kw, val, src) in enumerate(fillers):
                if i in used:
                    continue
                if kw in label:
                    set_cell_text(cells[1], str(val) if val is not None else "—",
                                  source=src)
                    used.add(i)
                    break

        if not tbl.rows:
            continue
        first = tbl.rows[0].cells
        # timeline table (header: 時間 / 事件 / 紀錄位置)
        if len(first) >= 3 and "時間" in cell_text(first[0]) and "事件" in cell_text(first[1]):
            fill_timeline_g2(tbl, yaml_data.get("timeline") or {})
        # checklist table
        if len(first) >= 3 and "項目" in cell_text(first[1]) and len(tbl.rows) >= 8:
            fill_checklist_g2(tbl, yaml_data.get("evidence_checklist") or {})

    fill_qa_g2(doc, yaml_data.get("likely_questions") or [])

    doc.save(str(output_path))
    return output_path


def format_euroscore_block(es):
    """Build a multi-line text block summarizing the EuroSCORE II result."""
    if not es or not es.get("computed"):
        return None, None
    comp = es["computed"]
    inp = es.get("inputs") or {}
    rat = es.get("rationale") or {}
    score = comp.get("score_pct")
    band_zh = {
        "low": "低風險",
        "intermediate": "中風險",
        "high": "高風險",
        "very_high": "極高風險",
    }.get(comp.get("risk_band"), comp.get("risk_band"))
    contribs = comp.get("contributors") or []
    top3 = "; ".join(c["factor"] for c in contribs[:3])

    lines = [f"EuroSCORE II = {score:.2f}%   ({band_zh})"]
    if comp.get("cc_ml_min") is not None:
        lines.append(f"  Cockcroft-Gault CC = {comp['cc_ml_min']:.1f} mL/min  → renal: {comp.get('derived_renal')}")
    if top3:
        lines.append(f"  主要貢獻因子: {top3}")
    lines.append("")
    lines.append("  Inputs:")
    lines.append(
        f"  age={inp.get('age')}  female={inp.get('female')}  "
        f"NYHA={inp.get('nyha')}  CCS4={inp.get('ccs4')}  "
        f"IDDM={inp.get('iddm')}  ECA={inp.get('extracardiac_arteriopathy')}  "
        f"COPD={inp.get('chronic_pulmonary_disease')}  poor mobility={inp.get('poor_mobility')}"
    )
    lines.append(
        f"  prev cardiac surgery={inp.get('previous_cardiac_surgery')}  "
        f"active IE={inp.get('active_endocarditis')}  critical preop={inp.get('critical_preop')}  "
        f"LV={inp.get('lv_function')}  recent MI={inp.get('recent_mi')}  "
        f"PASP={inp.get('pa_systolic')}  renal={inp.get('renal')}"
    )
    lines.append(
        f"  urgency={inp.get('urgency')}  weight={inp.get('weight_of_procedure')}  "
        f"thoracic aorta={inp.get('thoracic_aorta')}"
    )
    if rat:
        lines.append("")
        lines.append("  Rationale per field:")
        for k, v in rat.items():
            lines.append(f"    • {k}: {v}")

    src = "計算: euroscore_NCKUH/euroscore_ii.py (Nashef 2012)\n      Inputs 來源見上方各欄位 _src;院內 EuroSCORE 表單 (若有): 成大電子表單(依類別) → TemplateCode EMR-3-04-008"
    return "\n".join(lines), src


def insert_euroscore_after_why_pci(doc, euroscore_text, euroscore_src):
    """Find the table containing "為何選 PCI 不選 CABG" and insert a 2-col table
    right after it: [EuroSCORE II (預測手術死亡率) | <text>]."""
    if not euroscore_text:
        return
    target = None
    for tbl in doc.tables:
        joined = "".join(cell_text(c) for row in tbl.rows for c in row.cells)
        if "為何選 PCI" in joined:
            target = tbl
            break
    if target is None:
        return
    # Build a new table with the same column widths by cloning the target's structure.
    from copy import deepcopy
    new_tbl_xml = deepcopy(target._element)
    # We want exactly one row with 2 cells. Keep the first row, drop the rest.
    rows = new_tbl_xml.findall(qn("w:tr"))
    for r in rows[1:]:
        new_tbl_xml.remove(r)
    # Insert immediately after target
    target._element.addnext(new_tbl_xml)
    # Now grab the inserted table via the doc.tables list (rebuilt from XML)
    # The inserted table is right after `target` in document order.
    # Use python-docx's Table wrapper:
    from docx.table import Table
    new_table = Table(new_tbl_xml, target._parent)
    cells = new_table.rows[0].cells
    if len(cells) >= 2:
        set_cell_text(cells[0], "EuroSCORE II\n(預測手術死亡率)")
        set_cell_text(cells[1], euroscore_text, source=euroscore_src)


def insert_score_block_after_presentation(doc, hasbled, timi):
    """For ACS/AMI cases, insert HAS-BLED + TIMI block after Presentation row."""
    if not (hasbled or timi):
        return
    target = None
    for tbl in doc.tables:
        joined = "".join(cell_text(c) for row in tbl.rows for c in row.cells)
        if "Presentation" in joined and "STEMI" in joined:
            target = tbl
            break
    if target is None:
        return
    parts = []
    src_parts = []
    if hasbled:
        parts.append(f"HAS-BLED Score = {hasbled.get('total')} 分  ({hasbled.get('risk_label', '')})")
        items = hasbled.get("items") or []
        for it in items:
            parts.append(f"  {it}")
        if hasbled.get("src"):
            src_parts.append(hasbled["src"])
    if timi:
        if parts:
            parts.append("")
        parts.append(f"TIMI Score = {timi.get('total')} 分")
        items = timi.get("items") or []
        for it in items:
            parts.append(f"  {it}")
        if timi.get("src"):
            src_parts.append(timi["src"])
    text = "\n".join(parts)
    src = "; ".join(sorted(set(src_parts))) if src_parts else None

    from copy import deepcopy
    new_tbl_xml = deepcopy(target._element)
    rows = new_tbl_xml.findall(qn("w:tr"))
    for r in rows[1:]:
        new_tbl_xml.remove(r)
    target._element.addnext(new_tbl_xml)
    from docx.table import Table
    new_table = Table(new_tbl_xml, target._parent)
    cells = new_table.rows[0].cells
    if len(cells) >= 2:
        set_cell_text(cells[0], "HAS-BLED + TIMI\n(ACS 評分)")
        set_cell_text(cells[1], text, source=src)


def fill_doc(template_path, yaml_data, output_path):
    doc = Document(str(template_path))

    drop_group_two(doc)

    fillers = scalar_fillers(yaml_data)
    used = set()

    for tbl in doc.tables:
        for row in tbl.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            label = cell_text(cells[0]).strip()
            if not label:
                continue
            for i, (kw, val, src) in enumerate(fillers):
                if i in used:
                    continue
                if kw in label:
                    set_cell_text(cells[1], str(val) if val is not None else "—", source=src)
                    used.add(i)
                    break

        first_row_cells = tbl.rows[0].cells if tbl.rows else []
        if len(first_row_cells) >= 3:
            first_text = cell_text(first_row_cells[1])
            if "項目" in first_text:
                fill_checklist(tbl, yaml_data.get("evidence_checklist") or {})

    fill_qa(doc, yaml_data.get("likely_questions") or [])

    # NEW: insert score rows
    es = yaml_data.get("euroscore_ii") or {}
    es_text, es_src = format_euroscore_block(es)
    insert_euroscore_after_why_pci(doc, es_text, es_src)
    insert_score_block_after_presentation(
        doc,
        yaml_data.get("hasbled"),
        yaml_data.get("timi"),
    )

    doc.save(str(output_path))
    return output_path


def main():
    gemini_out = GEMINI_WORKSPACE / "out"
    for case in load_cases():
        chart = case["chart"]
        group = (case.get("group") or "1").strip()
        yaml_path = gemini_out / f"case_{chart}_filled.yaml"
        if not yaml_path.exists():
            print(f"  {chart}: skip ({yaml_path.name} not found)")
            continue
        y = yaml.safe_load(yaml_path.read_text(encoding="utf-8"))
        name_code = (y.get("name_code") or chart).replace("/", "")
        out_path = OUT_DIR / f"{chart}_{name_code}_filled.docx"
        try:
            if group == "2":
                fill_doc_g2(TEMPLATE, y, out_path)
            else:
                fill_doc(TEMPLATE, y, out_path)
            print(f"  rendered (group {group}) -> {out_path}")
        except PermissionError:
            print(f"  {chart}: SKIP — {out_path.name} is open (close it in Word and re-run)")


if __name__ == "__main__":
    main()
